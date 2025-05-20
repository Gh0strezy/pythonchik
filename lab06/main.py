import PySimpleGUI as sg
from docx import Document
import psycopg2
from psycopg2 import sql
from datetime import datetime

# Данные для расчета
data = {
    'Пиджак': {'44-46': 1.5, '48-50': 1.7, '52-54': 2.0, 'фурнитура': 300, 'работа': 1500},
    'Брюки': {'44-46': 1.2, '48-50': 1.4, '52-54': 1.6, 'фурнитура': 200, 'работа': 1000},
    'Костюм-тройка': {'44-46': 2.7, '48-50': 3.1, '52-54': 3.5, 'фурнитура': 500, 'работа': 2500}
}
price_per_meter = 500

sg.theme('LightGrey1')

layout = [
    [sg.Text('Тип изделия:', size=(15,1)), 
     sg.Radio('Пиджак', "TYPE", key='-JACKET-', default=True),
     sg.Radio('Брюки', "TYPE", key='-TROUSERS-'),
     sg.Radio('Костюм-тройка', "TYPE", key='-SUIT-')],
    
    [sg.Text('Размер:', size=(15,1)),
     sg.Radio('44-46', "SIZE", key='-44-46-', default=True),
     sg.Radio('48-50', "SIZE", key='-48-50-'),
     sg.Radio('52-54', "SIZE", key='-52-54-')],
    
    [sg.HorizontalSeparator()],
    
    [sg.Text('Расход ткани:', size=(15,1)), sg.InputText(size=(15,1), key='-TISSUE-', disabled=True)],
    [sg.Text('Стоимость ткани:', size=(15,1)), sg.InputText(size=(15,1), key='-TISSUE_COST-', disabled=True)],
    [sg.Text('Фурнитура:', size=(15,1)), sg.InputText(size=(15,1), key='-HARDWARE-', disabled=True)],
    [sg.Text('Работа:', size=(15,1)), sg.InputText(size=(15,1), key='-WORK-', disabled=True)],
    [sg.Text('Итого:', size=(15,1)), sg.InputText(size=(15,1), key='-TOTAL-', disabled=True)],
    
    [sg.Button('Рассчитать', size=(10,1)), sg.Button('Сохранить в Word', size=(15,1)), sg.Button('Сохранить в БД', size=(15,1)), sg.Exit('Выход', size=(10,1))]
]

window = sg.Window('Ателье - расчет стоимости', layout, finalize=True)

def get_selected_radio(values, group_keys):
    for key in group_keys:
        if values.get(key):
            return key
    return None

def calculate(values):
    # Определяем выбранный тип изделия
    clothing_key = get_selected_radio(values, ['-JACKET-', '-TROUSERS-', '-SUIT-'])
    clothing_map = {'-JACKET-': 'Пиджак', '-TROUSERS-': 'Брюки', '-SUIT-': 'Костюм-тройка'}
    clothing = clothing_map.get(clothing_key)
    
    # Определяем выбранный размер
    size_key = get_selected_radio(values, ['-44-46-', '-48-50-', '-52-54-'])
    size_map = {'-44-46-': '44-46', '-48-50-': '48-50', '-52-54-': '52-54'}
    size = size_map.get(size_key)
    
    if not clothing or not size:
        sg.popup_error('Пожалуйста, выберите изделие и размер')
        return None
    
    tissue = data[clothing][size]
    tissue_cost = tissue * price_per_meter
    hardware_cost = data[clothing]['фурнитура']
    work_cost = data[clothing]['работа']
    total = tissue_cost + hardware_cost + work_cost
    
    return {
        'clothing': clothing,
        'size': size,
        'tissue': tissue,
        'tissue_cost': tissue_cost,
        'hardware_cost': hardware_cost,
        'work_cost': work_cost,
        'total': total
    }

def save_to_word(report):
    filename = sg.popup_get_file('Сохранить отчет Word', save_as=True, no_window=True, file_types=(("Word files","*.docx"),))
    if not filename:
        return
    if not filename.endswith('.docx'):
        filename += '.docx'
    doc = Document()
    doc.add_heading('Отчет по расчету стоимости', 0)
    doc.add_paragraph(f"Тип изделия: {report['clothing']}")
    doc.add_paragraph(f"Размер: {report['size']}")
    doc.add_paragraph(f"Расход ткани: {report['tissue']:.1f} м")
    doc.add_paragraph(f"Стоимость ткани: {report['tissue_cost']:.0f} руб")
    doc.add_paragraph(f"Фурнитура: {report['hardware_cost']:.0f} руб")
    doc.add_paragraph(f"Работа: {report['work_cost']:.0f} руб")
    doc.add_paragraph(f"Итого: {report['total']:.0f} руб")
    doc.save(filename)
    sg.popup(f'Отчет сохранён в {filename}')

def save_to_db(report):
    try:
        conn = psycopg2.connect(
            dbname='atelier',
            user='postgres',
            password='ZXCZXC',
            host='localhost',
            port=5432
        )
        cur = conn.cursor()
        insert_query = psycopg2.sql.SQL("""
            INSERT INTO calculations 
            (clothing, size, tissue, tissue_cost, hardware_cost, work_cost, total, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """)
        cur.execute(insert_query, (
            report['clothing'], report['size'], report['tissue'], report['tissue_cost'],
            report['hardware_cost'], report['work_cost'], report['total'], datetime.now()
        ))
        conn.commit()
        cur.close()
        conn.close()
        sg.popup('Результаты успешно сохранены в базу данных')
    except Exception as e:
        sg.popup_error(f'Ошибка при сохранении в БД:\n{e}')


while True:
    event, values = window.read()
    if event in (sg.WIN_CLOSED, 'Выход'):
        break
    
    if event == 'Рассчитать':
        report = calculate(values)
        if report:
            window['-TISSUE-'].update(f"{report['tissue']:.1f} м")
            window['-TISSUE_COST-'].update(f"{report['tissue_cost']:.0f} руб")
            window['-HARDWARE-'].update(f"{report['hardware_cost']:.0f} руб")
            window['-WORK-'].update(f"{report['work_cost']:.0f} руб")
            window['-TOTAL-'].update(f"{report['total']:.0f} руб")
    
    elif event == 'Сохранить в Word':
        if not values['-TOTAL-']:
            sg.popup_error('Сначала выполните расчет')
            continue
        save_to_word(report)

    elif event == "Сохранить в БД":
        if not values['-TOTAL-']:
            sg.popup_error("Сначала выполните расчет")
            continue
        save_to_db(report)
window.close()
